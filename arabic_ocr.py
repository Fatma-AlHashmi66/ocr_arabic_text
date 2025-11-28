#!/usr/bin/env python3
"""
Arabic OCR PDF Extraction Tool

A comprehensive solution to extract Arabic text from PDF files containing images
using OCR and AI methods while preserving the exact document structure.

Usage:
    python arabic_ocr.py input.pdf [--output-dir OUTPUT_DIR] [--formats txt md docx]
    python arabic_ocr.py input_folder/ --batch [--output-dir OUTPUT_DIR]
"""

import argparse
import logging
import os
import sys
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional, Tuple, Union

import cv2
import numpy as np
from PIL import Image
from tqdm import tqdm

# Import configuration
try:
    from config import (
        BATCH_CONFIG,
        LOGGING_CONFIG,
        OCR_CONFIG,
        OUTPUT_CONFIG,
        PDF_CONFIG,
        PREPROCESSING_CONFIG,
        STRUCTURE_CONFIG,
    )
except ImportError:
    # Default configuration if config.py not found
    OCR_CONFIG = {"primary_engine": "tesseract", "min_confidence": 60}
    PREPROCESSING_CONFIG = {"enabled": True}
    PDF_CONFIG = {"dpi": 300}
    STRUCTURE_CONFIG = {"rtl": True, "page_separator": "\n--- Page {page_num} ---\n"}
    OUTPUT_CONFIG = {"formats": ["txt", "md", "docx"]}
    LOGGING_CONFIG = {"level": "INFO", "console": True}
    BATCH_CONFIG = {"continue_on_error": True, "show_progress": True}


@dataclass
class TextBlock:
    """Represents a block of text with position and content."""
    text: str
    x: int
    y: int
    width: int
    height: int
    confidence: float = 0.0
    line_num: int = 0
    block_num: int = 0


@dataclass
class PageResult:
    """Represents OCR results for a single page."""
    page_num: int
    text_blocks: List[TextBlock] = field(default_factory=list)
    raw_text: str = ""
    structured_text: str = ""
    confidence: float = 0.0
    image_path: Optional[str] = None


@dataclass
class DocumentResult:
    """Represents OCR results for an entire document."""
    source_path: str
    pages: List[PageResult] = field(default_factory=list)
    total_pages: int = 0
    average_confidence: float = 0.0


def setup_logging():
    """Configure logging based on settings."""
    level = getattr(logging, LOGGING_CONFIG.get("level", "INFO"))
    format_str = LOGGING_CONFIG.get(
        "format", "%(asctime)s - %(name)s - %(levelname)s - %(message)s"
    )
    
    handlers = []
    if LOGGING_CONFIG.get("console", True):
        handlers.append(logging.StreamHandler(sys.stdout))
    
    log_file = LOGGING_CONFIG.get("file")
    if log_file:
        handlers.append(logging.FileHandler(log_file, encoding="utf-8"))
    
    logging.basicConfig(level=level, format=format_str, handlers=handlers)
    return logging.getLogger("arabic_ocr")


logger = setup_logging()


class ImagePreprocessor:
    """Handles image preprocessing for better OCR accuracy."""
    
    def __init__(self, config: dict = None):
        self.config = config or PREPROCESSING_CONFIG
    
    def preprocess(self, image: np.ndarray) -> np.ndarray:
        """Apply all enabled preprocessing steps to the image."""
        if not self.config.get("enabled", True):
            return image
        
        result = image.copy()
        
        # Convert to grayscale if needed
        if len(result.shape) == 3:
            gray = cv2.cvtColor(result, cv2.COLOR_BGR2GRAY)
        else:
            gray = result
        
        # Deskewing
        if self.config.get("deskew", True):
            gray = self._deskew(gray)
        
        # Noise removal
        if self.config.get("denoise", True):
            strength = self.config.get("denoise_strength", 10)
            gray = cv2.fastNlMeansDenoising(gray, None, strength, 7, 21)
        
        # Contrast enhancement
        if self.config.get("enhance_contrast", True):
            gray = self._enhance_contrast(gray)
        
        # Binarization
        if self.config.get("binarize", False):
            threshold = self.config.get("binarize_threshold", 128)
            _, gray = cv2.threshold(gray, threshold, 255, cv2.THRESH_BINARY)
        
        return gray
    
    def _deskew(self, image: np.ndarray) -> np.ndarray:
        """Correct skew in the image."""
        # Find edges
        edges = cv2.Canny(image, 50, 150, apertureSize=3)
        
        # Detect lines using Hough transform
        lines = cv2.HoughLinesP(edges, 1, np.pi/180, 100, minLineLength=100, maxLineGap=10)
        
        if lines is None or len(lines) == 0:
            return image
        
        # Calculate average angle
        angles = []
        for line in lines:
            x1, y1, x2, y2 = line[0]
            if x2 != x1:
                angle = np.arctan2(y2 - y1, x2 - x1) * 180 / np.pi
                if abs(angle) < 45:  # Only consider near-horizontal lines
                    angles.append(angle)
        
        if not angles:
            return image
        
        median_angle = np.median(angles)
        
        # Rotate image to correct skew
        if abs(median_angle) > 0.5:  # Only deskew if angle is significant
            (h, w) = image.shape[:2]
            center = (w // 2, h // 2)
            rotation_matrix = cv2.getRotationMatrix2D(center, median_angle, 1.0)
            rotated = cv2.warpAffine(
                image, rotation_matrix, (w, h),
                flags=cv2.INTER_CUBIC,
                borderMode=cv2.BORDER_REPLICATE
            )
            return rotated
        
        return image
    
    def _enhance_contrast(self, image: np.ndarray) -> np.ndarray:
        """Enhance contrast using CLAHE."""
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
        return clahe.apply(image)


class PDFProcessor:
    """Handles PDF to image conversion."""
    
    def __init__(self, config: dict = None):
        self.config = config or PDF_CONFIG
    
    def extract_images(self, pdf_path: str) -> List[Tuple[int, Image.Image]]:
        """Extract images from PDF pages."""
        pdf_path = Path(pdf_path)
        if not pdf_path.exists():
            raise FileNotFoundError(f"PDF file not found: {pdf_path}")
        
        images = []
        dpi = self.config.get("dpi", 300)
        max_pages = self.config.get("max_pages")
        
        try:
            # Try using PyMuPDF (fitz) first
            import fitz
            doc = fitz.open(pdf_path)
            total_pages = len(doc)
            
            if max_pages:
                total_pages = min(total_pages, max_pages)
            
            for page_num in range(total_pages):
                page = doc[page_num]
                # Calculate zoom factor based on desired DPI
                zoom = dpi / 72  # 72 is default PDF resolution
                mat = fitz.Matrix(zoom, zoom)
                pix = page.get_pixmap(matrix=mat)
                
                # Convert to PIL Image
                img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
                
                if self.config.get("grayscale", True):
                    img = img.convert("L")
                
                images.append((page_num + 1, img))
            
            doc.close()
            logger.info(f"Extracted {len(images)} pages from PDF using PyMuPDF")
            
        except ImportError:
            # Fallback to pdf2image
            try:
                from pdf2image import convert_from_path
                
                pil_images = convert_from_path(
                    pdf_path,
                    dpi=dpi,
                    grayscale=self.config.get("grayscale", True),
                    thread_count=self.config.get("thread_count", -1),
                    last_page=max_pages,
                )
                
                images = [(i + 1, img) for i, img in enumerate(pil_images)]
                logger.info(f"Extracted {len(images)} pages from PDF using pdf2image")
                
            except ImportError:
                raise ImportError(
                    "Neither PyMuPDF nor pdf2image is installed. "
                    "Please install one: pip install PyMuPDF or pip install pdf2image"
                )
        
        return images


class TesseractOCR:
    """Tesseract OCR engine wrapper."""
    
    def __init__(self, config: dict = None):
        self.config = config or OCR_CONFIG.get("tesseract", {})
        try:
            import pytesseract
            self.pytesseract = pytesseract
        except ImportError:
            raise ImportError("pytesseract not installed. Run: pip install pytesseract")
    
    def recognize(self, image: Union[Image.Image, np.ndarray]) -> List[TextBlock]:
        """Perform OCR on an image and return text blocks with positions."""
        if isinstance(image, np.ndarray):
            image = Image.fromarray(image)
        
        lang = self.config.get("lang", "ara")
        psm = self.config.get("psm", 3)
        oem = self.config.get("oem", 1)
        custom_config = self.config.get("config", "")
        
        config_str = f"--psm {psm} --oem {oem} {custom_config}"
        
        # Get detailed OCR data
        try:
            data = self.pytesseract.image_to_data(
                image, lang=lang, config=config_str, output_type=self.pytesseract.Output.DICT
            )
        except Exception as e:
            logger.error(f"Tesseract OCR failed: {e}")
            return []
        
        text_blocks = []
        n_boxes = len(data["text"])
        
        for i in range(n_boxes):
            text = data["text"][i].strip()
            if text:
                conf = float(data["conf"][i])
                if conf >= 0:  # -1 indicates no text
                    block = TextBlock(
                        text=text,
                        x=data["left"][i],
                        y=data["top"][i],
                        width=data["width"][i],
                        height=data["height"][i],
                        confidence=conf,
                        line_num=data["line_num"][i],
                        block_num=data["block_num"][i],
                    )
                    text_blocks.append(block)
        
        return text_blocks
    
    def recognize_simple(self, image: Union[Image.Image, np.ndarray]) -> str:
        """Perform simple OCR returning just the text."""
        if isinstance(image, np.ndarray):
            image = Image.fromarray(image)
        
        lang = self.config.get("lang", "ara")
        psm = self.config.get("psm", 3)
        oem = self.config.get("oem", 1)
        custom_config = self.config.get("config", "")
        
        config_str = f"--psm {psm} --oem {oem} {custom_config}"
        
        try:
            text = self.pytesseract.image_to_string(image, lang=lang, config=config_str)
            return text
        except Exception as e:
            logger.error(f"Tesseract simple OCR failed: {e}")
            return ""


class EasyOCREngine:
    """EasyOCR engine wrapper."""
    
    def __init__(self, config: dict = None):
        self.config = config or OCR_CONFIG.get("easyocr", {})
        self.reader = None
    
    def _get_reader(self):
        """Lazy initialization of EasyOCR reader."""
        if self.reader is None:
            try:
                import easyocr
                languages = self.config.get("languages", ["ar", "en"])
                gpu = self.config.get("gpu", False)
                self.reader = easyocr.Reader(languages, gpu=gpu)
            except ImportError:
                raise ImportError("easyocr not installed. Run: pip install easyocr")
        return self.reader
    
    def recognize(self, image: Union[Image.Image, np.ndarray]) -> List[TextBlock]:
        """Perform OCR on an image and return text blocks with positions."""
        if isinstance(image, Image.Image):
            image = np.array(image)
        
        reader = self._get_reader()
        
        text_threshold = self.config.get("text_threshold", 0.7)
        link_threshold = self.config.get("link_threshold", 0.4)
        low_text = self.config.get("low_text", 0.4)
        
        try:
            results = reader.readtext(
                image,
                text_threshold=text_threshold,
                link_threshold=link_threshold,
                low_text=low_text,
            )
        except Exception as e:
            logger.error(f"EasyOCR failed: {e}")
            return []
        
        text_blocks = []
        for i, (bbox, text, conf) in enumerate(results):
            # bbox is [[x1,y1], [x2,y2], [x3,y3], [x4,y4]]
            x_coords = [p[0] for p in bbox]
            y_coords = [p[1] for p in bbox]
            x = int(min(x_coords))
            y = int(min(y_coords))
            width = int(max(x_coords) - x)
            height = int(max(y_coords) - y)
            
            block = TextBlock(
                text=text.strip(),
                x=x,
                y=y,
                width=width,
                height=height,
                confidence=conf * 100,  # Convert to percentage
                line_num=i,
                block_num=0,
            )
            text_blocks.append(block)
        
        return text_blocks


class StructurePreserver:
    """Preserves document structure from OCR results."""
    
    def __init__(self, config: dict = None):
        self.config = config or STRUCTURE_CONFIG
    
    def reconstruct_text(self, text_blocks: List[TextBlock], image_height: int = 0) -> str:
        """Reconstruct text from blocks preserving structure."""
        if not text_blocks:
            return ""
        
        # Sort blocks by position (top to bottom, right to left for RTL)
        rtl = self.config.get("rtl", True)
        
        # Group blocks by line
        lines = self._group_into_lines(text_blocks)
        
        # Sort lines by vertical position
        lines = sorted(lines, key=lambda blocks: min(b.y for b in blocks))
        
        result_lines = []
        prev_y = 0
        
        for line_blocks in lines:
            # Calculate vertical spacing
            current_y = min(b.y for b in line_blocks)
            if prev_y > 0:
                gap = current_y - prev_y
                avg_height = np.mean([b.height for b in line_blocks])
                if gap > avg_height * self.config.get("vertical_space_threshold", 1.5):
                    # Add extra blank lines for large gaps
                    extra_lines = int(gap / avg_height) - 1
                    result_lines.extend([""] * min(extra_lines, 3))
            
            # Sort blocks within line (RTL or LTR)
            if rtl:
                line_blocks = sorted(line_blocks, key=lambda b: -b.x)
            else:
                line_blocks = sorted(line_blocks, key=lambda b: b.x)
            
            # Reconstruct line with spacing
            line_text = self._reconstruct_line(line_blocks, rtl)
            result_lines.append(line_text)
            
            prev_y = current_y + max(b.height for b in line_blocks)
        
        return "\n".join(result_lines)
    
    def _group_into_lines(self, blocks: List[TextBlock]) -> List[List[TextBlock]]:
        """Group text blocks that appear on the same line."""
        if not blocks:
            return []
        
        # Sort by y position
        sorted_blocks = sorted(blocks, key=lambda b: b.y)
        
        lines = []
        current_line = [sorted_blocks[0]]
        current_y = sorted_blocks[0].y
        
        for block in sorted_blocks[1:]:
            # Check if block is on the same line (within height tolerance)
            avg_height = np.mean([b.height for b in current_line])
            if abs(block.y - current_y) < avg_height * 0.5:
                current_line.append(block)
            else:
                lines.append(current_line)
                current_line = [block]
                current_y = block.y
        
        if current_line:
            lines.append(current_line)
        
        return lines
    
    def _reconstruct_line(self, blocks: List[TextBlock], rtl: bool) -> str:
        """Reconstruct a single line with appropriate spacing."""
        if not blocks:
            return ""
        
        if len(blocks) == 1:
            return blocks[0].text
        
        result_parts = []
        threshold = self.config.get("horizontal_space_threshold", 2.0)
        
        for i, block in enumerate(blocks):
            result_parts.append(block.text)
            
            if i < len(blocks) - 1:
                next_block = blocks[i + 1]
                if rtl:
                    gap = block.x - (next_block.x + next_block.width)
                else:
                    gap = next_block.x - (block.x + block.width)
                
                avg_width = np.mean([b.width / max(1, len(b.text)) for b in blocks])
                
                if gap > avg_width * threshold:
                    # Add extra space for large gaps
                    extra_spaces = min(int(gap / avg_width), 5)
                    result_parts.append(" " * extra_spaces)
                else:
                    result_parts.append(" ")
        
        return "".join(result_parts)


class OutputGenerator:
    """Generates output in various formats."""
    
    def __init__(self, config: dict = None):
        self.config = config or OUTPUT_CONFIG
    
    def save_text(self, result: DocumentResult, output_path: Path) -> None:
        """Save result as plain text file."""
        txt_config = self.config.get("txt", {})
        encoding = txt_config.get("encoding", "utf-8")
        include_pages = txt_config.get("include_page_numbers", True)
        
        with open(output_path, "w", encoding=encoding) as f:
            for page in result.pages:
                if include_pages:
                    f.write(f"\n{'=' * 50}\n")
                    f.write(f"Page {page.page_num}\n")
                    f.write(f"{'=' * 50}\n\n")
                
                f.write(page.structured_text or page.raw_text)
                f.write("\n")
        
        logger.info(f"Saved text file: {output_path}")
    
    def save_markdown(self, result: DocumentResult, output_path: Path) -> None:
        """Save result as Markdown file."""
        md_config = self.config.get("md", {})
        page_break = md_config.get("page_break_style", "hr")
        
        with open(output_path, "w", encoding="utf-8") as f:
            # Title
            source_name = Path(result.source_path).stem
            f.write(f"# {source_name}\n\n")
            f.write(f"> Extracted using Arabic OCR\n")
            f.write(f"> Total pages: {result.total_pages}\n")
            f.write(f"> Average confidence: {result.average_confidence:.1f}%\n\n")
            
            for page in result.pages:
                if page_break == "hr":
                    f.write("---\n\n")
                else:
                    f.write(f"## Page {page.page_num}\n\n")
                
                # Wrap text in RTL div for proper rendering
                f.write('<div dir="rtl">\n\n')
                
                text = page.structured_text or page.raw_text
                # Convert line breaks to markdown
                text = text.replace("\n\n", "\n\n")
                f.write(text)
                
                f.write("\n\n</div>\n\n")
        
        logger.info(f"Saved Markdown file: {output_path}")
    
    def save_docx(self, result: DocumentResult, output_path: Path) -> None:
        """Save result as Word document."""
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        except ImportError:
            logger.error("python-docx not installed. Run: pip install python-docx")
            return
        
        docx_config = self.config.get("docx", {})
        font_name = docx_config.get("font_name", "Arial")
        font_size = docx_config.get("font_size", 12)
        rtl = docx_config.get("rtl_direction", True)
        page_breaks = docx_config.get("include_page_breaks", True)
        
        doc = Document()
        
        # Set document properties for RTL
        for section in doc.sections:
            section.page_width = Pt(595)  # A4 width
            section.page_height = Pt(842)  # A4 height
        
        for i, page in enumerate(result.pages):
            text = page.structured_text or page.raw_text
            
            # Add page header
            header = doc.add_paragraph()
            header_run = header.add_run(f"Page {page.page_num}")
            header_run.bold = True
            header_run.font.size = Pt(font_size + 2)
            header_run.font.name = font_name
            
            # Add text content
            for line in text.split("\n"):
                para = doc.add_paragraph()
                run = para.add_run(line)
                run.font.name = font_name
                run.font.size = Pt(font_size)
                
                if rtl:
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            
            # Add page break
            if page_breaks and i < len(result.pages) - 1:
                doc.add_page_break()
        
        doc.save(output_path)
        logger.info(f"Saved Word document: {output_path}")


class ArabicOCR:
    """Main Arabic OCR processing class."""
    
    def __init__(
        self,
        ocr_engine: str = None,
        use_fallback: bool = None,
        min_confidence: float = None,
    ):
        self.ocr_engine = ocr_engine or OCR_CONFIG.get("primary_engine", "tesseract")
        self.use_fallback = use_fallback if use_fallback is not None else OCR_CONFIG.get("use_fallback", True)
        self.min_confidence = min_confidence or OCR_CONFIG.get("min_confidence", 60)
        
        self.preprocessor = ImagePreprocessor()
        self.pdf_processor = PDFProcessor()
        self.structure_preserver = StructurePreserver()
        self.output_generator = OutputGenerator()
        
        # Initialize OCR engines
        self.tesseract = None
        self.easyocr = None
        
        if self.ocr_engine == "tesseract":
            self.tesseract = TesseractOCR()
            if self.use_fallback:
                try:
                    self.easyocr = EasyOCREngine()
                except ImportError:
                    logger.warning("EasyOCR not available as fallback")
        else:
            self.easyocr = EasyOCREngine()
            if self.use_fallback:
                try:
                    self.tesseract = TesseractOCR()
                except ImportError:
                    logger.warning("Tesseract not available as fallback")
    
    def process_pdf(
        self,
        pdf_path: str,
        output_dir: str = None,
        formats: List[str] = None,
    ) -> DocumentResult:
        """Process a PDF file and extract Arabic text."""
        pdf_path = Path(pdf_path)
        logger.info(f"Processing PDF: {pdf_path}")
        
        # Create output directory
        if output_dir:
            output_dir = Path(output_dir)
        else:
            output_dir = pdf_path.parent / "output"
        output_dir.mkdir(parents=True, exist_ok=True)
        
        # Extract images from PDF
        images = self.pdf_processor.extract_images(str(pdf_path))
        
        result = DocumentResult(
            source_path=str(pdf_path),
            total_pages=len(images),
        )
        
        # Process each page
        show_progress = BATCH_CONFIG.get("show_progress", True)
        iterator = tqdm(images, desc="Processing pages") if show_progress else images
        
        total_confidence = 0
        for page_num, image in iterator:
            page_result = self._process_page(page_num, image)
            result.pages.append(page_result)
            total_confidence += page_result.confidence
        
        result.average_confidence = total_confidence / len(images) if images else 0
        
        # Save outputs
        formats = formats or OUTPUT_CONFIG.get("formats", ["txt", "md", "docx"])
        base_name = pdf_path.stem
        
        for fmt in formats:
            output_path = output_dir / f"{base_name}.{fmt}"
            if fmt == "txt":
                self.output_generator.save_text(result, output_path)
            elif fmt == "md":
                self.output_generator.save_markdown(result, output_path)
            elif fmt == "docx":
                self.output_generator.save_docx(result, output_path)
        
        logger.info(
            f"Completed processing {pdf_path.name}: "
            f"{result.total_pages} pages, "
            f"average confidence: {result.average_confidence:.1f}%"
        )
        
        return result
    
    def _process_page(self, page_num: int, image: Image.Image) -> PageResult:
        """Process a single page image."""
        # Convert to numpy array for preprocessing
        np_image = np.array(image)
        
        # Preprocess image
        preprocessed = self.preprocessor.preprocess(np_image)
        
        # Perform OCR
        text_blocks, confidence = self._perform_ocr(preprocessed)
        
        # Reconstruct structured text
        structured_text = self.structure_preserver.reconstruct_text(
            text_blocks, image_height=preprocessed.shape[0]
        )
        
        # Get raw text as fallback
        raw_text = " ".join(block.text for block in text_blocks)
        
        return PageResult(
            page_num=page_num,
            text_blocks=text_blocks,
            raw_text=raw_text,
            structured_text=structured_text,
            confidence=confidence,
        )
    
    def _perform_ocr(self, image: np.ndarray) -> Tuple[List[TextBlock], float]:
        """Perform OCR using configured engine(s)."""
        text_blocks = []
        confidence = 0
        
        # Try primary engine
        primary_engine = self.tesseract if self.ocr_engine == "tesseract" else self.easyocr
        
        if primary_engine:
            try:
                text_blocks = primary_engine.recognize(image)
                if text_blocks:
                    confidence = np.mean([b.confidence for b in text_blocks])
            except Exception as e:
                logger.error(f"Primary OCR engine failed: {e}")
        
        # Try fallback if confidence is low
        if self.use_fallback and confidence < self.min_confidence:
            fallback_engine = self.easyocr if self.ocr_engine == "tesseract" else self.tesseract
            
            if fallback_engine:
                try:
                    fallback_blocks = fallback_engine.recognize(image)
                    if fallback_blocks:
                        fallback_confidence = np.mean([b.confidence for b in fallback_blocks])
                        if fallback_confidence > confidence:
                            text_blocks = fallback_blocks
                            confidence = fallback_confidence
                            logger.debug(
                                f"Switched to fallback engine "
                                f"(confidence: {fallback_confidence:.1f}%)"
                            )
                except Exception as e:
                    logger.error(f"Fallback OCR engine failed: {e}")
        
        return text_blocks, confidence
    
    def process_batch(
        self,
        input_path: str,
        output_dir: str = None,
        formats: List[str] = None,
    ) -> List[DocumentResult]:
        """Process multiple PDF files."""
        input_path = Path(input_path)
        
        if input_path.is_file():
            pdf_files = [input_path]
        else:
            pdf_files = list(input_path.glob("*.pdf"))
        
        if not pdf_files:
            logger.warning(f"No PDF files found in {input_path}")
            return []
        
        logger.info(f"Processing {len(pdf_files)} PDF files")
        
        results = []
        continue_on_error = BATCH_CONFIG.get("continue_on_error", True)
        
        for pdf_file in tqdm(pdf_files, desc="Processing PDFs"):
            try:
                result = self.process_pdf(str(pdf_file), output_dir, formats)
                results.append(result)
            except Exception as e:
                logger.error(f"Failed to process {pdf_file}: {e}")
                if not continue_on_error:
                    raise
        
        return results


def create_arg_parser() -> argparse.ArgumentParser:
    """Create command line argument parser."""
    parser = argparse.ArgumentParser(
        description="Extract Arabic text from PDF files using OCR",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Process a single PDF
  python arabic_ocr.py document.pdf
  
  # Process with specific output formats
  python arabic_ocr.py document.pdf --formats txt md
  
  # Process all PDFs in a directory
  python arabic_ocr.py ./pdfs/ --batch
  
  # Use EasyOCR instead of Tesseract
  python arabic_ocr.py document.pdf --engine easyocr
        """,
    )
    
    parser.add_argument(
        "input",
        help="Input PDF file or directory (with --batch)",
    )
    
    parser.add_argument(
        "-o", "--output-dir",
        help="Output directory for extracted text",
    )
    
    parser.add_argument(
        "-f", "--formats",
        nargs="+",
        choices=["txt", "md", "docx"],
        default=["txt", "md", "docx"],
        help="Output formats (default: txt md docx)",
    )
    
    parser.add_argument(
        "-e", "--engine",
        choices=["tesseract", "easyocr"],
        default="tesseract",
        help="OCR engine to use (default: tesseract)",
    )
    
    parser.add_argument(
        "--batch",
        action="store_true",
        help="Process all PDF files in the input directory",
    )
    
    parser.add_argument(
        "--no-fallback",
        action="store_true",
        help="Disable fallback to secondary OCR engine",
    )
    
    parser.add_argument(
        "--min-confidence",
        type=float,
        default=60,
        help="Minimum confidence threshold (0-100, default: 60)",
    )
    
    parser.add_argument(
        "-v", "--verbose",
        action="store_true",
        help="Enable verbose output",
    )
    
    return parser


def main():
    """Main entry point."""
    parser = create_arg_parser()
    args = parser.parse_args()
    
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)
    
    ocr = ArabicOCR(
        ocr_engine=args.engine,
        use_fallback=not args.no_fallback,
        min_confidence=args.min_confidence,
    )
    
    if args.batch:
        results = ocr.process_batch(
            args.input,
            output_dir=args.output_dir,
            formats=args.formats,
        )
        logger.info(f"Processed {len(results)} files")
    else:
        result = ocr.process_pdf(
            args.input,
            output_dir=args.output_dir,
            formats=args.formats,
        )
        logger.info(
            f"Processed {result.total_pages} pages "
            f"with average confidence {result.average_confidence:.1f}%"
        )


if __name__ == "__main__":
    main()
